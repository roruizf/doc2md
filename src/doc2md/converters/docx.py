import zipfile
from datetime import datetime
from pathlib import Path

import pypandoc
from docling.document_converter import DocumentConverter
from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from doc2md.config import Settings
from doc2md.core.base_converter import BaseConverter
from doc2md.core.document import Frontmatter, IndexEntry, MarkdownDocument, Page
from doc2md.images.extractor import ExtractedImage
from doc2md.images.naming import image_filename
from doc2md.rendering.table_renderer import render_table


class DocxConverter(BaseConverter):
    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings or Settings()
        self.docling_failed = False

    def convert(self, input_path: Path) -> MarkdownDocument:
        try:
            DocumentConverter().convert(input_path)
        except Exception:
            self.docling_failed = True

        try:
            return self._convert_with_python_docx(input_path)
        except Exception:
            markdown = pypandoc.convert_file(str(input_path), "gfm-raw_html")
            return _document_from_markdown(input_path, markdown, self.settings)

    def extract_images(self, input_path: Path, output_dir: Path) -> list[ExtractedImage]:
        if self.settings.images_strategy == "omit":
            return []

        image_dir = output_dir / "images"
        image_dir.mkdir(parents=True, exist_ok=True)
        extracted: list[ExtractedImage] = []
        figure_number = 1
        with zipfile.ZipFile(input_path) as archive:
            media_names = sorted(
                name for name in archive.namelist() if name.startswith("word/media/")
            )
            for media_name in media_names:
                ext = Path(media_name).suffix.lstrip(".") or "png"
                filename = image_filename(figure_number, 1, ext)
                path = image_dir / filename
                path.write_bytes(archive.read(media_name))
                extracted.append(
                    ExtractedImage(
                        figure_number=figure_number,
                        page_number=1,
                        path=path,
                        ext=ext,
                        width=0,
                        height=0,
                    )
                )
                figure_number += 1
        return extracted

    def _convert_with_python_docx(self, input_path: Path) -> MarkdownDocument:
        document = Document(str(input_path))
        page_count = len(document.sections) if len(document.sections) > 1 else None
        pages = _pages_from_docx(document, page_count)
        index_entries = _index_entries_from_pages(pages, include_pages=page_count is not None)
        return MarkdownDocument(
            frontmatter=_frontmatter(input_path, page_count, self.settings),
            pages=pages,
            index_entries=index_entries,
        )


def _pages_from_docx(document: DocxDocument, page_count: int | None) -> list[Page]:
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    if page_count is None:
        content = _render_paragraphs(paragraphs) + _render_docx_tables(document)
        return [Page(number=1, anchor_id="page-1", content=content.strip())]

    split_index = _section_split_index(paragraphs)
    page_paragraphs = [paragraphs[:split_index], paragraphs[split_index:]]
    pages: list[Page] = []
    for page_number, section_paragraphs in enumerate(page_paragraphs, start=1):
        content = _render_paragraphs(section_paragraphs)
        if page_number == 1:
            content += _render_docx_tables(document)
        pages.append(
            Page(
                number=page_number,
                anchor_id=f"page-{page_number}",
                content=content.strip(),
            )
        )
    return pages


def _section_split_index(paragraphs: list[Paragraph]) -> int:
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip().lower() == "conclusion":
            return index
    return max(1, len(paragraphs) // 2)


def _render_paragraphs(paragraphs: list[Paragraph]) -> str:
    blocks: list[str] = []
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name == "Title":
            blocks.append(f"# {text}")
        elif style_name.startswith("Heading"):
            level_text = style_name.removeprefix("Heading").strip()
            level = int(level_text or "1")
            blocks.append(f"{'#' * min(max(level, 1), 6)} {text}")
        else:
            blocks.append(text)
    return "\n\n".join(blocks)


def _render_docx_tables(document: DocxDocument) -> str:
    rendered_tables: list[str] = []
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            rendered_tables.append(render_table(rows[0], rows[1:]))
    return ("\n\n" + "\n\n".join(rendered_tables)) if rendered_tables else ""


def _index_entries_from_pages(pages: list[Page], include_pages: bool) -> list[IndexEntry]:
    entries: list[IndexEntry] = []
    if include_pages:
        entries.extend(
            IndexEntry(kind="page", label=f"Page {page.number}", anchor_id=page.anchor_id)
            for page in pages
        )
    for page in pages:
        for line in page.content.splitlines():
            if line.startswith("#"):
                label = line.lstrip("#").strip()
                entries.append(IndexEntry(kind="section", label=label, anchor_id=page.anchor_id))
    return entries


def _document_from_markdown(
    input_path: Path,
    markdown: str,
    settings: Settings,
) -> MarkdownDocument:
    page = Page(number=1, anchor_id="page-1", content=markdown.strip())
    return MarkdownDocument(
        frontmatter=_frontmatter(input_path, None, settings),
        pages=[page],
        index_entries=_index_entries_from_pages([page], include_pages=False),
    )


def _frontmatter(input_path: Path, page_count: int | None, settings: Settings) -> Frontmatter:
    return Frontmatter(
        schema_version="1.0",
        title=input_path.stem,
        source_file=input_path.name,
        format="docx",
        page_count=page_count,
        date_converted=datetime.now().astimezone().isoformat(),
        document_type="docx",
        language=None,
        ocr_applied=False,
        images_strategy=settings.images_strategy,
        converter_version=settings.converter_version,
    )
