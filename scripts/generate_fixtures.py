#!/usr/bin/env python3
"""Generate reproducible test fixtures for doc2md.

Install fixture-generation dependencies with:

    pip install reportlab pymupdf pikepdf pillow python-docx ebooklib pypandoc
"""

from __future__ import annotations

import shutil
import tempfile
from pathlib import Path
from typing import Callable

import fitz
import pikepdf
import pypandoc
from docx import Document
from docx.enum.section import WD_SECTION
from ebooklib import epub
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as ReportLabImage,
    PageBreak,
    PageTemplate,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
FIXTURES = ROOT / "tests" / "fixtures"

EXPECTED_FIXTURES = [
    "sample_digital.pdf",
    "sample_multicolumn.pdf",
    "sample_scanned.pdf",
    "sample_locked.pdf",
    "sample_mixed.pdf",
    "sample.docx",
    "sample.odt",
    "sample.epub",
    "sample.html",
    "sample.txt",
    "sample_image.png",
]

CLASSIC_TEXT = (
    "It is a truth universally acknowledged, that a single man in possession "
    "of a good fortune, must be in want of a wife. However little known the "
    "feelings or views of such a man may be on his first entering a "
    "neighbourhood, this truth is so well fixed in the minds of the surrounding "
    "families, that he is considered as the rightful property of some one or "
    "other of their daughters."
)


def _fixture_path(filename: str) -> Path:
    return FIXTURES / filename


def _log_generated(filename: str) -> None:
    print(f"Generated: {filename}")


def _simple_png(path: Path, size: tuple[int, int] = (100, 100), color: str = "red") -> None:
    image = Image.new("RGB", size, color)
    draw = ImageDraw.Draw(image)
    draw.rectangle((10, 10, size[0] - 10, size[1] - 10), outline="black", width=3)
    image.save(path)


def generate_sample_image() -> None:
    path = _fixture_path("sample_image.png")
    image = Image.new("RGB", (600, 400), "white")
    draw = ImageDraw.Draw(image)
    text = "Hello world from doc2md"
    font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), text, font=font)
    x = (image.width - (bbox[2] - bbox[0])) // 2
    y = (image.height - (bbox[3] - bbox[1])) // 2
    draw.text((x, y), text, fill="black", font=font)
    image.save(path)
    _log_generated(path.name)


def generate_digital_pdf() -> None:
    path = _fixture_path("sample_digital.pdf")
    temp_image = _fixture_path("_embedded_red_square.png")
    _simple_png(temp_image)

    styles = getSampleStyleSheet()
    story = [
        Paragraph("Sample Digital PDF", styles["Title"]),
        Paragraph("Introduction", styles["Heading1"]),
        Paragraph(CLASSIC_TEXT, styles["BodyText"]),
        Spacer(1, 12),
        Paragraph("This first page is intentionally simple digital text.", styles["BodyText"]),
        PageBreak(),
        Paragraph("Data Table", styles["Heading1"]),
        Paragraph("A compact table for extraction tests.", styles["BodyText"]),
        Table(
            [
                ["Quarter", "Revenue", "Growth", "Notes"],
                ["Q1", "12.4", "+5.2%", "Baseline"],
                ["Q2", "14.1", "+13.7%", "Expansion"],
                ["Q3", "15.8", "+12.1%", "Stable"],
                ["Q4", "17.2", "+8.9%", "Close"],
            ],
            hAlign="LEFT",
        ),
        PageBreak(),
        Paragraph("Figure Page", styles["Heading1"]),
        Paragraph("This page includes one embedded PNG image.", styles["BodyText"]),
        ReportLabImage(str(temp_image), width=100, height=100),
        Paragraph("The image above is a red square fixture.", styles["BodyText"]),
    ]

    for item in story:
        if isinstance(item, Table):
            item.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ]
                )
            )

    SimpleDocTemplate(str(path), pagesize=letter).build(story)
    temp_image.unlink(missing_ok=True)
    _log_generated(path.name)


def generate_multicolumn_pdf() -> None:
    path = _fixture_path("sample_multicolumn.pdf")
    styles = getSampleStyleSheet()
    doc = BaseDocTemplate(str(path), pagesize=letter)
    width, height = letter
    margin = 54
    gap = 18
    column_width = (width - (2 * margin) - gap) / 2
    frames = [
        Frame(margin, margin, column_width, height - (2 * margin), id="left"),
        Frame(margin + column_width + gap, margin, column_width, height - (2 * margin), id="right"),
    ]
    doc.addPageTemplates([PageTemplate(id="two-column", frames=frames)])

    story = [Paragraph("Two Column Fixture", styles["Title"])]
    for idx in range(18):
        prefix = "LEFT COLUMN MARKER" if idx == 0 else "RIGHT COLUMN MARKER" if idx == 9 else "Column text"
        story.append(Paragraph(f"{prefix} {idx + 1}. {CLASSIC_TEXT}", styles["BodyText"]))
        story.append(Spacer(1, 8))
    story.append(PageBreak())
    story.append(Paragraph("Second Page", styles["Heading1"]))
    for idx in range(8):
        story.append(Paragraph(f"Continuation paragraph {idx + 1}. {CLASSIC_TEXT}", styles["BodyText"]))
        story.append(Spacer(1, 8))
    doc.build(story)
    _log_generated(path.name)


def generate_scanned_pdf() -> None:
    source = _fixture_path("sample_digital.pdf")
    path = _fixture_path("sample_scanned.pdf")
    scanned = fitz.open()
    source_doc = fitz.open(source)
    for page in source_doc:
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
        rect = page.rect
        new_page = scanned.new_page(width=rect.width, height=rect.height)
        new_page.insert_image(rect, stream=pix.tobytes("png"))
    scanned.save(path)
    scanned.close()
    source_doc.close()
    _log_generated(path.name)


def generate_locked_pdf() -> None:
    path = _fixture_path("sample_locked.pdf")
    source = _fixture_path("sample_digital.pdf")
    with pikepdf.open(source) as pdf:
        pdf.save(
            path,
            encryption=pikepdf.Encryption(
                user="test123",
                owner="owner456",
                R=6,
            ),
        )
    _log_generated(path.name)


def generate_mixed_pdf() -> None:
    path = _fixture_path("sample_mixed.pdf")
    digital = pikepdf.open(_fixture_path("sample_digital.pdf"))
    scanned = pikepdf.open(_fixture_path("sample_scanned.pdf"))
    output = pikepdf.Pdf.new()
    output.pages.append(digital.pages[0])
    output.pages.append(digital.pages[1])
    output.pages.append(scanned.pages[0])
    output.save(path)
    output.close()
    digital.close()
    scanned.close()
    _log_generated(path.name)


def generate_docx() -> None:
    path = _fixture_path("sample.docx")
    image_path = _fixture_path("_docx_red_square.png")
    _simple_png(image_path)

    doc = Document()
    doc.add_heading("Sample Document", 0)
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This DOCX fixture contains headings, paragraphs, a table, and an image.")
    doc.add_heading("Data Table", level=2)
    table = doc.add_table(rows=3, cols=3)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.text = f"R{row_idx + 1}C{col_idx + 1}"
    doc.add_heading("Section Two", level=1)
    doc.add_paragraph("This paragraph appears before the section break.")
    doc.add_picture(str(image_path), width=None)
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("The second section closes the sample document.")
    doc.save(path)
    image_path.unlink(missing_ok=True)
    _log_generated(path.name)


def generate_odt() -> None:
    path = _fixture_path("sample.odt")
    pypandoc.convert_file(str(_fixture_path("sample.docx")), "odt", outputfile=str(path))
    _log_generated(path.name)


def generate_epub() -> None:
    path = _fixture_path("sample.epub")
    image_path = _fixture_path("_epub_red_square.png")
    _simple_png(image_path)

    book = epub.EpubBook()
    book.set_identifier("doc2md-sample-epub")
    book.set_title("Sample EPUB")
    book.set_language("en")
    book.add_author("doc2md Test")

    image_item = epub.EpubItem(
        uid="red_square",
        file_name="images/red_square.png",
        media_type="image/png",
        content=image_path.read_bytes(),
    )
    book.add_item(image_item)

    chapters = []
    chapter_specs = [
        ("intro", "Chapter 1: Introduction", "<p>This is the introduction chapter.</p>"),
        (
            "data",
            "Chapter 2: Data",
            '<p>This chapter has an inline image.</p><img src="images/red_square.png" alt="red square"/>',
        ),
        ("conclusion", "Chapter 3: Conclusion", "<p>This is the conclusion chapter.</p>"),
    ]
    for uid, title, body in chapter_specs:
        chapter = epub.EpubHtml(title=title, file_name=f"{uid}.xhtml", lang="en")
        chapter.content = f"<html><body><h1>{title}</h1>{body}</body></html>"
        book.add_item(chapter)
        chapters.append(chapter)

    book.toc = tuple(chapters)
    book.spine = ["nav", *chapters]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(path), book)
    image_path.unlink(missing_ok=True)
    _log_generated(path.name)


def generate_html() -> None:
    path = _fixture_path("sample.html")
    path.write_text(
        """<!doctype html>
<html lang="en">
<head><meta charset="utf-8"><title>Sample HTML</title></head>
<body>
  <h1>Sample HTML Document</h1>
  <h2>Subheading</h2>
  <p>This HTML fixture includes a paragraph, table, and image reference.</p>
  <table>
    <thead><tr><th>Name</th><th>Value</th><th>Note</th></tr></thead>
    <tbody>
      <tr><td>Alpha</td><td>1</td><td>First</td></tr>
      <tr><td>Beta</td><td>2</td><td>Second</td></tr>
    </tbody>
  </table>
  <img src="./sample_image.png" alt="Sample image">
</body>
</html>
""",
        encoding="utf-8",
    )
    _log_generated(path.name)


def generate_txt() -> None:
    path = _fixture_path("sample.txt")
    path.write_text(
        """INTRODUCTION
============

This text fixture contains simple prose and a cross-reference: see page 12, Table 3.

BACKGROUND
----------

The second section uses an underline marker and more paragraph text.

CONCLUSION
==========

Final observations are intentionally short and deterministic.
""",
        encoding="utf-8",
    )
    _log_generated(path.name)


def _run_step(filename: str, func: Callable[[], None]) -> bool:
    try:
        func()
    except Exception as exc:  # noqa: BLE001 - continue fixture generation after failures.
        print(f"Error generating {filename}: {exc}")
        return False
    return _fixture_path(filename).exists()


def main() -> int:
    FIXTURES.mkdir(parents=True, exist_ok=True)

    for filename in EXPECTED_FIXTURES:
        _fixture_path(filename).unlink(missing_ok=True)

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        # Keep tempfile imported and exercised for compatibility with the requested script shape.
        shutil.rmtree(tmp_path, ignore_errors=True)

    steps: list[tuple[str, Callable[[], None]]] = [
        ("sample_image.png", generate_sample_image),
        ("sample_digital.pdf", generate_digital_pdf),
        ("sample_multicolumn.pdf", generate_multicolumn_pdf),
        ("sample_scanned.pdf", generate_scanned_pdf),
        ("sample_locked.pdf", generate_locked_pdf),
        ("sample_mixed.pdf", generate_mixed_pdf),
        ("sample.docx", generate_docx),
        ("sample.odt", generate_odt),
        ("sample.epub", generate_epub),
        ("sample.html", generate_html),
        ("sample.txt", generate_txt),
    ]

    successes = 0
    for filename, func in steps:
        if _run_step(filename, func):
            successes += 1

    if successes == len(EXPECTED_FIXTURES):
        print("All 11 fixtures ready.")
    print(f"{successes}/11 fixtures generated successfully")
    return 0 if successes == len(EXPECTED_FIXTURES) else 1


if __name__ == "__main__":
    raise SystemExit(main())
